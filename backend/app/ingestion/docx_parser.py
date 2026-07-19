"""
DOCX text extraction — Phase 3 (DOCX slice).

Uses python-docx directly — no OCR involved, .docx files always have a real
text layer (they're XML under the hood, never scanned images). Extracts
both paragraph text and table cell text, since policy/SOP documents often
put key details (e.g. PPE requirement matrices, maintenance intervals) in
tables rather than plain paragraphs.

Returns the same ParsedDocument/ParsedPage shape as pdf_parser.py and
ocr_parser.py so pipeline.py can treat all three extraction paths identically
downstream. DOCX has no real concept of "pages" outside of print layout, so
the whole file is returned as a single ParsedPage (page_number=1).
"""

from docx import Document as DocxDocument

from app.ingestion.pdf_parser import ParsedDocument, ParsedPage


def _extract_table_text(table) -> list[str]:
    """Flattens a docx table into row-by-row text lines, cells joined by ' | '
    so structured data (e.g. a PPE-by-zone matrix) stays readable as plain
    text rather than losing its row/column relationships entirely."""
    lines = []
    for row in table.rows:
        cells_text = [cell.text.strip() for cell in row.cells]
        if any(cells_text):
            lines.append(" | ".join(cells_text))
    return lines


def parse_docx(filepath) -> ParsedDocument:
    """Extracts all paragraph and table text from a .docx file, in document
    order. Raises ValueError if the result is suspiciously empty — a
    near-blank docx is worth flagging rather than silently storing nothing."""
    doc = DocxDocument(str(filepath))

    text_blocks = []

    # python-docx exposes paragraphs and tables as separate top-level
    # collections, not interleaved in document order — good enough for the
    # MVP to just emit all paragraphs, then all tables, rather than
    # reconstructing exact document order (which needs body.iter_inner_content()
    # walking, a heavier lift not needed for retrieval-quality text).
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_blocks.append(text)

    for table in doc.tables:
        table_lines = _extract_table_text(table)
        text_blocks.extend(table_lines)

    full_text = "\n".join(text_blocks)

    if len(full_text.strip()) < 20:
        raise ValueError(
            f"{filepath}: extracted almost no text from this .docx file. "
            f"Check it isn't a near-empty template or corrupted file."
        )

    page = ParsedPage(page_number=1, text=full_text)
    return ParsedDocument(
        filename=str(filepath),
        page_count=1,
        full_text=full_text,
        pages=[page],
        extractor_used="python-docx",
    )


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 2:
        print("Usage: python -m app.ingestion.docx_parser <path-to-docx>")
        sys.exit(1)

    parsed = parse_docx(sys.argv[1])
    print(f"Extractor used : {parsed.extractor_used}")
    print(f"Total chars    : {len(parsed.full_text)}")
    print("\n--- First 500 characters ---\n")
    print(parsed.full_text[:500])
